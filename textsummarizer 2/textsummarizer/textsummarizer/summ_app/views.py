
# # from django.shortcuts import render
# # from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
# # import torch
# # import re

# # MODEL_NAME = "facebook/bart-large-cnn"  # You can also try: "sshleifer/distilbart-cnn-12-6"
# # tokenizer = AutoTokenizer.from_pretrained(MODEL_NAME)
# # model = AutoModelForSeq2SeqLM.from_pretrained(MODEL_NAME)

# # def clean_input(text):
# #     return re.sub(r'\s+', ' ', text.strip())

# # def summarize_text(text):
# #     cleaned_text = clean_input(text)

# #     inputs = tokenizer.encode(
# #         cleaned_text,
# #         return_tensors="pt",
# #         max_length=1024,
# #         truncation=True
# #     )

# #     # Adjust generation parameters to improve summary quality
# #     summary_ids = model.generate(
# #         inputs,
# #         max_length=150,
# #         min_length=40,
# #         length_penalty=1.5,
# #         num_beams=5,
# #         early_stopping=True,
# #         no_repeat_ngram_size=3,
# #         repetition_penalty=2.0
# #     )

# #     summary = tokenizer.decode(summary_ids[0], skip_special_tokens=True)

# #     # Split into clean bullet points
# #     points = re.split(r'[.!?]', summary)
# #     return [point.strip() for point in points if point.strip()]

# # def summarize_view(request):
# #     summary = None
# #     error = None

# #     if request.method == "POST":
# #         input_text = request.POST.get("text", "").strip()

# #         if not input_text:
# #             error = "Please enter some text to summarize."
# #         else:
# #             try:
# #                 summary = summarize_text(input_text)
# #             except Exception as e:
# #                 error = f"Error: {str(e)}"

# #     return render(request, "summarizer/summarize.html", {
# #         "summary": summary,
# #         "error": error
# #     })




# from django.shortcuts import render
# from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
# import torch
# import re

# # Load model and tokenizer
# MODEL_NAME = "facebook/bart-large-cnn"
# tokenizer = AutoTokenizer.from_pretrained(MODEL_NAME)
# model = AutoModelForSeq2SeqLM.from_pretrained(MODEL_NAME)

# def clean_input(text):
#     return re.sub(r'\s+', ' ', text.strip())

# def summarize_text(text):
#     cleaned_text = clean_input(text)

#     inputs = tokenizer.encode(
#         cleaned_text,
#         return_tensors="pt",
#         max_length=1024,
#         truncation=True
#     )
#     # print("Input token length:", len(inputs[0]))


#     summary_ids = model.generate(
#         inputs,
#         max_length=300,
#         min_length=80,
#         length_penalty=1.5,
#         num_beams=4,
#         early_stopping=True,
#         no_repeat_ngram_size=3,
#         repetition_penalty=2.0
#     )

#     summary = tokenizer.decode(summary_ids[0], skip_special_tokens=True)

#     # Custom formatting: split summary into meaningful chunks for point-wise display
#     summary = re.sub(r'\s*\.\s*', '.\n', summary)  # Insert line breaks after periods
#     lines = [line.strip() for line in summary.split('\n') if line.strip()]
#     return lines

# def summarize_view(request):
#     summary = None
#     error = None

#     if request.method == "POST":
#         input_text = request.POST.get("text", "").strip()

#         if not input_text:
#             error = "Please enter some text to summarize."
#         else:
#             try:
#                 summary = summarize_text(input_text)
#             except Exception as e:
#                 error = f"Error: {str(e)}"

#     return render(request, "summarizer/summarize.html", {
#         "summary": summary,
#         "error": error
#     })







from django.shortcuts import render
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
import torch
import re
from docx import Document  # Add this

# Load model and tokenizer
MODEL_NAME = "facebook/bart-large-cnn"
tokenizer = AutoTokenizer.from_pretrained(MODEL_NAME)
model = AutoModelForSeq2SeqLM.from_pretrained(MODEL_NAME)

def clean_input(text):
    return re.sub(r'\s+', ' ', text.strip())

def summarize_text(text):
    cleaned_text = clean_input(text)

    inputs = tokenizer.encode(
        cleaned_text,
        return_tensors="pt",
        max_length=1024,
        truncation=True
    )

    summary_ids = model.generate(
        inputs,
        max_length=400,    # Pehle 300 tha, ab 400 tak allow karo
        min_length=120,    # Pehle 80 tha, ab 120 ka minimum maintain karo
        length_penalty=1.2,   # Thoda relaxed
        num_beams=6,      # Beams thoda zyada karo better results ke liye
        early_stopping=True,
        no_repeat_ngram_size=3,
        repetition_penalty=1.5
    )

    summary = tokenizer.decode(summary_ids[0], skip_special_tokens=True)

    # Line break ke liye
    summary = re.sub(r'\s*\.\s*', '.\n', summary)
    lines = [line.strip() for line in summary.split('\n') if line.strip()]
    return lines

def extract_text_from_docx(file):
    """Extract text from a Word (.docx) file."""
    doc = Document(file)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)





def summarize_view(request):
    summary = None
    error = None

    if request.method == "POST":
        input_text = request.POST.get("text", "").strip()

        if not input_text and 'docx_file' in request.FILES:
            docx_file = request.FILES['docx_file']
            if docx_file.name.endswith(".docx"):
                from docx import Document
                document = Document(docx_file)
                input_text = "\n".join([para.text for para in document.paragraphs if para.text.strip()])
        
        if not input_text:
            error = "Please enter some text or upload a Word document."
        else:
            try:
                summary_points = summarize_text(input_text)
                # extended_summary = extend_summary(summary_points)  # << extended summary
                summary = summary_points
            except Exception as e:
                error = f"Error: {str(e)}"

    return render(request, "summarizer/summarize.html", {
        "summary": summary,
        "error": error
    })




